from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_word_report(json_data, filename=None):
    doc = Document()
    
    # --- 1. TITLE & HEADER ---
    title = doc.add_heading('GigGuard | Incident Memory Log', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    timestamp = doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp.style = "Subtitle"
    
    doc.add_paragraph("_" * 70) 

    # --- 2. METADATA ---
    meta = json_data.get('meta', {})
    p = doc.add_paragraph()
    p.add_run("Report Type: ").bold = True
    p.add_run(f"{meta.get('report_type', 'Standard Report')}\n")
    p.add_run("Reference ID: ").bold = True
    p.add_run(f"{meta.get('report_id', 'N/A')}")

    # --- 3. CLASSIFICATION ---
    doc.add_heading('1. Classification', level=1)
    
    classification = json_data.get('classification', {})
    severity_level = str(classification.get('severity_level') or 'Medium')

    severity_paragraph = doc.add_paragraph()
    run_label = severity_paragraph.add_run("SEVERITY LEVEL: ")
    run_label.bold = True
    
    run_sev = severity_paragraph.add_run(severity_level.upper())
    run_sev.bold = True
    run_sev.font.size = Pt(14)
    
    if severity_level in ["High", "Critical"]:
        run_sev.font.color.rgb = RGBColor(220, 0, 0)
    elif severity_level == "Medium":
        run_sev.font.color.rgb = RGBColor(255, 140, 0)
    else:
        run_sev.font.color.rgb = RGBColor(0, 150, 0)

    p = doc.add_paragraph()
    p.add_run("Category: ").bold = True
    p.add_run(f"{classification.get('primary_category', 'Uncategorized')}\n")
    p.add_run("Keywords: ").bold = True
    keywords = classification.get('keywords', [])
    p.add_run(", ".join(keywords) if keywords else "None")

    # --- 4. NARRATIVE ---
    doc.add_heading('2. Incident Narrative', level=1)
    narrative = json_data.get('narrative', {})
    doc.add_paragraph(str(narrative.get('objective_summary') or 'No summary available.'))

    # --- 5. TIMELINE (FIXED CRASH HERE) ---
    doc.add_heading('3. Chronological Timeline', level=1)
    
    timeline_data = narrative.get('chronological_timeline', [])
    if timeline_data:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light List Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Time Reference'
        hdr_cells[1].text = 'Event Description'
        
        for event in timeline_data:
            row_cells = table.add_row().cells
            
            # --- CRITICAL FIX START ---
            # Check if 'event' is a Dictionary (Correct) or String (AI Hallucination)
            if isinstance(event, dict):
                time_val = event.get('time_reference')
                desc_val = event.get('event')
            else:
                # If it's just a string, put the whole thing in Description
                time_val = "-"
                desc_val = str(event)
            # --------------------------

            row_cells[0].text = str(time_val or '-')
            row_cells[1].text = str(desc_val or '-')
    else:
        doc.add_paragraph("No timeline events detected.")

    # --- 6. ENTITIES ---
    doc.add_heading('4. Identified Entities', level=1)
    entities = json_data.get('entities', {})
    
    def add_entity_section(label, items):
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        if items:
            # Safe list conversion
            safe_items = [str(i) for i in items] if isinstance(items, list) else [str(items)]
            p.add_run(", ".join(safe_items))
        else:
            p.add_run("None Identified").italic = True

    add_entity_section("People", entities.get('people_involved', []))
    add_entity_section("Vehicles", entities.get('vehicles', []))
    add_entity_section("Damage/Injuries", entities.get('injuries_or_damages', []))

    # --- 7. LOCATION CONTEXT ---
    doc.add_heading('5. Location Verification', level=1)
    
    location_context = json_data.get('location_context', {})
    
    loc_table = doc.add_table(rows=2, cols=2)
    loc_table.style = 'Table Grid'
    
    row1 = loc_table.rows[0].cells
    row1[0].text = "System Recorded GPS:"
    row1[0].paragraphs[0].runs[0].bold = True
    row1[1].text = str(location_context.get('system_recorded_gps') or 'N/A')
    
    row2 = loc_table.rows[1].cells
    row2[0].text = "Mentioned in Audio:"
    row2[0].paragraphs[0].runs[0].bold = True
    row2[1].text = str(location_context.get('transcript_mentioned_location') or 'N/A')

    # --- 8. DISCLAIMER ---
    doc.add_paragraph("\n")
    disclaimer = doc.add_paragraph("DISCLAIMER: This document is an automated archival record based on user testimony. It does not constitute a verified legal finding.")
    disclaimer.style = "Quote" 
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if filename:
        doc.save(filename)
        print(f"Word Document saved locally as: {filename}")
        
    return doc